from __future__ import annotations

import csv
from collections import Counter
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Rust_Category_Detection_Model_Report.docx"

MODEL_RUNS = [
    {
        "name": "YOLOv8n OBB",
        "base_model": "yolov8n-obb.pt",
        "run_dir": ROOT / "runs" / "obb" / "runs" / "obb" / "rust_corrosion_yolov8_obb",
        "best_weights": ROOT
        / "runs"
        / "obb"
        / "runs"
        / "obb"
        / "rust_corrosion_yolov8_obb"
        / "weights"
        / "best.pt",
    },
    {
        "name": "YOLO11s OBB",
        "base_model": "yolo11s-obb.pt",
        "run_dir": ROOT / "runs" / "obb_yolo11" / "rust_corrosion_yolo11s_obb",
        "best_weights": ROOT
        / "runs"
        / "obb_yolo11"
        / "rust_corrosion_yolo11s_obb"
        / "weights"
        / "best.pt",
    },
]

CLASS_NAMES = {
    0: "mild-corrosion",
    1: "moderate-corrosion",
    2: "severe-corrosion",
}

IMAGE_ORDER = {
    "results.png": 0,
    "BoxPR_curve.png": 1,
    "BoxF1_curve.png": 2,
    "BoxP_curve.png": 3,
    "BoxR_curve.png": 4,
    "confusion_matrix.png": 5,
    "confusion_matrix_normalized.png": 6,
    "labels.jpg": 7,
}


def read_results(csv_path: Path) -> list[dict[str, str]]:
    with csv_path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def as_float(row: dict[str, str], key: str) -> float:
    return float(row[key.strip()] if key.strip() in row else row[key])


def metric(row: dict[str, str], name: str) -> float:
    normalized = {key.strip(): value for key, value in row.items()}
    return float(normalized[name])


def best_rows(rows: list[dict[str, str]]) -> dict[str, dict[str, str]]:
    return {
        "Best mAP50-95": max(rows, key=lambda row: metric(row, "metrics/mAP50-95(B)")),
        "Best mAP50": max(rows, key=lambda row: metric(row, "metrics/mAP50(B)")),
        "Best precision": max(rows, key=lambda row: metric(row, "metrics/precision(B)")),
        "Best recall": max(rows, key=lambda row: metric(row, "metrics/recall(B)")),
        "Final epoch": rows[-1],
    }


def read_args(args_path: Path) -> dict[str, str]:
    args: dict[str, str] = {}
    if not args_path.exists():
        return args
    for line in args_path.read_text(encoding="utf-8").splitlines():
        if ":" not in line or line.startswith(" "):
            continue
        key, value = line.split(":", 1)
        args[key.strip()] = value.strip()
    return args


def count_split(split: str) -> tuple[int, int, Counter[int]]:
    image_dir = ROOT / "data" / split / "images"
    label_dir = ROOT / "data" / split / "labels"
    image_count = sum(1 for _ in image_dir.glob("*")) if image_dir.exists() else 0
    label_files = list(label_dir.glob("*.txt")) if label_dir.exists() else []
    objects: Counter[int] = Counter()
    for label_file in label_files:
        for line in label_file.read_text(encoding="utf-8").splitlines():
            parts = line.split()
            if parts and parts[0].isdigit():
                objects[int(parts[0])] += 1
    return image_count, len(label_files), objects


def image_description(path: Path, model_name: str | None = None) -> str:
    name = path.name
    prefix = f"For {model_name}, this image" if model_name else "This image"

    if name == "results.png":
        return (
            f"{prefix} summarizes the training history across epochs, including box, "
            "classification, DFL, and angle losses together with precision, recall, "
            "mAP50, and mAP50-95 validation metrics."
        )
    if name == "BoxPR_curve.png":
        return (
            f"{prefix} shows the precision-recall curve for oriented rust detections. "
            "The area and curve shape indicate how well the model balances correct "
            "detections against missed detections across confidence thresholds."
        )
    if name == "BoxF1_curve.png":
        return (
            f"{prefix} shows F1 score across confidence thresholds. It highlights the "
            "threshold region where precision and recall are best balanced."
        )
    if name == "BoxP_curve.png":
        return (
            f"{prefix} shows precision across confidence thresholds. Higher precision "
            "means fewer false rust detections are produced."
        )
    if name == "BoxR_curve.png":
        return (
            f"{prefix} shows recall across confidence thresholds. Higher recall means "
            "more annotated rust regions are found by the model."
        )
    if name == "confusion_matrix.png":
        return (
            f"{prefix} is the raw confusion matrix, showing how detections are assigned "
            "to mild, moderate, and severe corrosion classes and where class confusion occurs."
        )
    if name == "confusion_matrix_normalized.png":
        return (
            f"{prefix} is the normalized confusion matrix, making class-wise accuracy "
            "and misclassification patterns easier to compare."
        )
    if name == "labels.jpg":
        return (
            f"{prefix} visualizes the dataset label distribution and oriented box "
            "placement patterns used during training."
        )
    if name.startswith("train_batch"):
        return (
            f"{prefix} shows a training batch with augmented input images and oriented "
            "ground-truth rust boxes used to teach the model."
        )
    if name.startswith("val_batch") and name.endswith("_labels.jpg"):
        return (
            f"{prefix} shows validation images with ground-truth oriented rust labels. "
            "These annotations are the reference used to evaluate predictions."
        )
    if name.startswith("val_batch") and name.endswith("_pred.jpg"):
        return (
            f"{prefix} shows validation predictions generated by the trained model, "
            "including oriented boxes and predicted corrosion categories."
        )
    if path.parts[-3:-1] == ("static", "results"):
        return (
            "This platform result image shows the web application output after running "
            "rust category detection on an uploaded inspection image."
        )
    if path.parts[-3:-1] == ("static", "samples"):
        return (
            "This platform sample image is used in the web interface to demonstrate "
            "rust inspection outputs and model performance visuals."
        )
    return "This image is a project artifact included for visual reference."


def sorted_images(paths: Iterable[Path]) -> list[Path]:
    return sorted(
        paths,
        key=lambda path: (
            IMAGE_ORDER.get(path.name, 20),
            0 if path.name.startswith("train_batch") else 1,
            path.name,
        ),
    )


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_key_value_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Value"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value


def add_metrics_table(document: Document, rows_by_label: dict[str, dict[str, str]]) -> None:
    headers = [
        "Metric point",
        "Epoch",
        "Precision",
        "Recall",
        "mAP50",
        "mAP50-95",
        "Box loss",
        "Class loss",
        "DFL loss",
        "Angle loss",
    ]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for label, row in rows_by_label.items():
        cells = table.add_row().cells
        values = [
            label,
            str(int(float(metric(row, "epoch")))),
            f"{metric(row, 'metrics/precision(B)'):.4f}",
            f"{metric(row, 'metrics/recall(B)'):.4f}",
            f"{metric(row, 'metrics/mAP50(B)'):.4f}",
            f"{metric(row, 'metrics/mAP50-95(B)'):.4f}",
            f"{metric(row, 'val/box_loss'):.4f}",
            f"{metric(row, 'val/cls_loss'):.4f}",
            f"{metric(row, 'val/dfl_loss'):.4f}",
            f"{metric(row, 'val/angle_loss'):.4f}",
        ]
        for idx, value in enumerate(values):
            cells[idx].text = value


def add_image(document: Document, path: Path, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(str(path), width=Inches(6.2))
    except Exception as exc:
        document.add_paragraph(f"Image could not be embedded: {path} ({exc})")
        return
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(f"{path.name}: {caption}")
    caption_run.italic = True


def main() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    document.add_heading("Rust Category Detection Platform - Model Training Report", 0)
    document.add_paragraph(
        "This document summarizes the achieved training metrics and visual artifacts "
        "for the two oriented object detection models used in the rust category "
        "detection platform. The models detect mild, moderate, and severe corrosion "
        "with oriented bounding boxes."
    )

    add_heading(document, "Dataset Summary")
    split_rows = []
    total_images = total_labels = total_objects = 0
    class_totals: Counter[int] = Counter()
    for split in ("train", "valid", "test"):
        images, labels, objects = count_split(split)
        total_images += images
        total_labels += labels
        total_objects += sum(objects.values())
        class_totals.update(objects)
        split_rows.append(
            (
                split,
                str(images),
                str(labels),
                str(objects.get(0, 0)),
                str(objects.get(1, 0)),
                str(objects.get(2, 0)),
                str(sum(objects.values())),
            )
        )

    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    for idx, header in enumerate(
        ["Split", "Images", "Label files", "Mild", "Moderate", "Severe", "Objects"]
    ):
        table.rows[0].cells[idx].text = header
    for row in split_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    document.add_paragraph(
        f"Total images: {total_images}; total label files: {total_labels}; total "
        f"annotated rust instances: {total_objects}. Class names: "
        + ", ".join(CLASS_NAMES.values())
        + "."
    )

    for model in MODEL_RUNS:
        document.add_section(WD_SECTION.NEW_PAGE)
        add_heading(document, model["name"])
        run_dir = model["run_dir"]
        args = read_args(run_dir / "args.yaml")
        rows = read_results(run_dir / "results.csv")
        achieved = best_rows(rows)

        add_key_value_table(
            document,
            [
                ("Base model", model["base_model"]),
                ("Run folder", str(run_dir.relative_to(ROOT))),
                ("Best weights", str(model["best_weights"].relative_to(ROOT))),
                ("Epochs completed", str(len(rows))),
                ("Image size", args.get("imgsz", "Not recorded")),
                ("Batch size", args.get("batch", "Not recorded")),
                ("Device", args.get("device", "Not recorded")),
                ("Optimizer", args.get("optimizer", "Not recorded")),
                ("Learning-rate schedule", "Cosine LR" if args.get("cos_lr") == "true" else "Standard LR"),
                ("Patience", args.get("patience", "Not recorded")),
            ],
        )

        document.add_paragraph(
            "The table below reports the strongest achieved validation points from "
            "the training CSV, plus the final epoch. mAP50-95 is the strictest summary "
            "metric because it averages performance across multiple IoU thresholds."
        )
        add_metrics_table(document, achieved)

        add_heading(document, "Training and Validation Images", level=2)
        for image_path in sorted_images(
            list(run_dir.glob("*.png")) + list(run_dir.glob("*.jpg"))
        ):
            add_image(document, image_path, image_description(image_path, model["name"]))

    platform_images = sorted_images(
        list((ROOT / "rust_detection_platform" / "static" / "samples").glob("*.*"))
        + list((ROOT / "rust_detection_platform" / "static" / "results").glob("*.*"))
    )
    if platform_images:
        document.add_section(WD_SECTION.NEW_PAGE)
        add_heading(document, "Platform Sample and Result Images")
        document.add_paragraph(
            "These images are included from the platform's static sample/result folders "
            "to show how the trained rust detection output appears in the application."
        )
        for image_path in platform_images:
            if image_path.suffix.lower() in {".jpg", ".jpeg", ".png"}:
                add_image(document, image_path, image_description(image_path))

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
